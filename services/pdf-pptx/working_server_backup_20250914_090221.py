from flask import Flask, request, render_template, send_file, jsonify
import os
import tempfile
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt
from docx.shared import RGBColor
import pytesseract
import cv2
import numpy as np
from PIL import Image

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

# Tesseract 경로 설정 (Windows)
if os.name == 'nt':
    tesseract_path = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
    if os.path.exists(tesseract_path):
        pytesseract.pytesseract.tesseract_cmd = tesseract_path

def safe_filename(filename, max_length=100):
    """안전한 파일명 생성"""
    if not filename:
        return "untitled"
    
    # 기본 보안 처리
    filename = secure_filename(filename)
    if not filename:
        return "untitled"
    
    # 길이 제한
    if len(filename) > max_length:
        name, ext = os.path.splitext(filename)
        filename = name[:max_length-len(ext)] + ext
    
    return filename

def ocr_image_to_blocks(image):
    """이미지에서 OCR로 텍스트 블록 추출"""
    try:
        # 이미지 전처리
        gray = cv2.cvtColor(np.array(image), cv2.COLOR_RGB2GRAY)
        denoised = cv2.fastNlMeansDenoising(gray)
        clahe = cv2.createCLAHE(clipLimit=2.0, tileGridSize=(8,8))
        enhanced = clahe.apply(denoised)
        
        # OCR 수행
        config = r"--oem 3 --psm 6 -l kor+eng"
        data = pytesseract.image_to_data(enhanced, config=config, output_type=pytesseract.Output.DICT)
        
        blocks = []
        n = len(data["text"])
        for i in range(n):
            text = data["text"][i].strip()
            conf_val = data["conf"][i]
            
            # conf 값 안전하게 처리
            if isinstance(conf_val, (int, float)):
                conf = int(conf_val)
            elif isinstance(conf_val, str) and conf_val.replace('.', '').replace('-', '').isdigit():
                conf = int(float(conf_val))
            else:
                conf = -1
                
            if text and conf >= 30:
                blocks.append({
                    "text": text,
                    "x": data["left"][i],
                    "y": data["top"][i],
                    "w": data["width"][i],
                    "h": data["height"][i],
                    "conf": conf,
                })
        
        return blocks
    except Exception as e:
        print(f"OCR 블록 추출 오류: {e}")
        return []

def pdf_to_docx_text_only(pdf_path, output_path):
    """PDF를 텍스트만 추출하여 DOCX로 변환"""
    try:
        print(f"📄 PDF → DOCX 변환 시작: {pdf_path}")
        
        # PDF를 이미지로 변환
        images = convert_from_path(pdf_path, dpi=200)
        print(f"📄 {len(images)} 페이지 변환 시작")
        
        # 새 DOCX 문서 생성
        doc = Document()
        
        for i, image in enumerate(images):
            print(f"  📄 페이지 {i+1}/{len(images)} 처리 중...")
            
            # OCR로 텍스트 추출
            blocks = ocr_image_to_blocks(image)
            
            if blocks:
                print(f"    ✅ {len(blocks)}개 텍스트 블록 발견")
                
                # Y 좌표 기준으로 정렬하여 줄별로 그룹화
                blocks.sort(key=lambda b: (b["y"], b["x"]))
                lines = []
                current_line = []
                line_y = -1
                
                for block in blocks:
                    if line_y < 0 or abs(block["y"] - line_y) < 20:  # 20픽셀 이내면 같은 줄
                        current_line.append(block)
                        line_y = block["y"]
                    else:
                        if current_line:
                            lines.append(current_line)
                        current_line = [block]
                        line_y = block["y"]
                
                if current_line:
                    lines.append(current_line)
                
                # 페이지 구분 추가 (첫 페이지가 아닌 경우)
                if i > 0:
                    doc.add_page_break()
                
                # 각 줄을 편집 가능한 문단으로 추가
                for line in lines:
                    words = sorted(line, key=lambda x: x["x"])
                    text = " ".join([w["text"] for w in words])
                    if text.strip():
                        p = doc.add_paragraph()
                        run = p.add_run(text)
                        run.font.name = "맑은 고딕"
                        run.font.size = Pt(12)
                        
                        # 평균 신뢰도에 따른 색상 조절
                        avg_conf = sum(w["conf"] for w in words) / len(words)
                        if avg_conf < 60:
                            run.font.color.rgb = RGBColor(100, 100, 100)  # 낮은 신뢰도는 회색
                        else:
                            run.font.color.rgb = RGBColor(0, 0, 0)  # 높은 신뢰도는 검은색
                        
                        p.paragraph_format.space_after = Pt(6)
            else:
                print("    ⚠️ 텍스트를 찾을 수 없음")
                # 페이지 구분 추가 (첫 페이지가 아닌 경우)
                if i > 0:
                    doc.add_page_break()
                
                # 빈 페이지 안내
                p = doc.add_paragraph()
                run = p.add_run(f"페이지 {i+1}: 텍스트를 감지할 수 없습니다. 여기에 직접 입력하세요.")
                run.font.name = "맑은 고딕"
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(150, 150, 150)
        
        # 문서 저장
        doc.save(output_path)
        print(f"✅ DOCX 변환 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ PDF → DOCX 변환 오류: {e}")
        return False

@app.route('/')
def index():
    return '''
    <!DOCTYPE html>
    <html>
    <head>
        <title>PDF 텍스트 변환기</title>
        <meta charset="utf-8">
        <style>
            body { font-family: Arial, sans-serif; max-width: 800px; margin: 50px auto; padding: 20px; }
            .upload-box { border: 2px dashed #ccc; padding: 40px; text-align: center; margin: 20px 0; }
            .btn { background: #007bff; color: white; padding: 10px 20px; border: none; border-radius: 5px; cursor: pointer; }
            .btn:hover { background: #0056b3; }
            .result { margin-top: 20px; padding: 20px; background: #f8f9fa; border-radius: 5px; }
        </style>
    </head>
    <body>
        <h1>📄 PDF → 편집 가능한 텍스트 DOCX 변환기</h1>
        <p><strong>새로운 기능:</strong> 원본 이미지 없이 추출된 텍스트만으로 편집 가능한 DOCX 문서를 생성합니다!</p>
        
        <div class="upload-box">
            <form id="uploadForm" enctype="multipart/form-data">
                <input type="file" id="fileInput" accept=".pdf" style="display: none;">
                <button type="button" onclick="document.getElementById('fileInput').click()" class="btn">PDF 파일 선택</button>
                <div id="fileName" style="margin-top: 10px; color: #666;"></div>
                <button type="submit" class="btn" style="margin-top: 10px; display: none;" id="convertBtn">텍스트 변환 시작</button>
            </form>
        </div>
        
        <div id="result" class="result" style="display: none;">
            <h3>변환 결과</h3>
            <div id="resultContent"></div>
        </div>
        
        <script>
            document.getElementById('fileInput').addEventListener('change', function(e) {
                const file = e.target.files[0];
                if (file) {
                    document.getElementById('fileName').textContent = '선택된 파일: ' + file.name;
                    document.getElementById('convertBtn').style.display = 'inline-block';
                }
            });
            
            document.getElementById('uploadForm').addEventListener('submit', function(e) {
                e.preventDefault();
                const fileInput = document.getElementById('fileInput');
                const file = fileInput.files[0];
                
                if (!file) {
                    alert('파일을 선택해주세요.');
                    return;
                }
                
                const formData = new FormData();
                formData.append('file', file);
                
                const resultDiv = document.getElementById('result');
                const resultContent = document.getElementById('resultContent');
                resultDiv.style.display = 'block';
                resultContent.innerHTML = '<p>📄 텍스트 변환 중... 잠시만 기다려주세요.</p>';
                
                fetch('/convert', {
                    method: 'POST',
                    body: formData
                })
                .then(response => response.json())
                .then(data => {
                    if (data.success) {
                        resultContent.innerHTML = `
                            <p>✅ 변환 완료!</p>
                            <a href="/download/${data.filename}" class="btn">📄 DOCX 파일 다운로드</a>
                        `;
                    } else {
                        resultContent.innerHTML = `<p>❌ 오류: ${data.error}</p>`;
                    }
                })
                .catch(error => {
                    resultContent.innerHTML = `<p>❌ 네트워크 오류: ${error}</p>`;
                });
            });
        </script>
    </body>
    </html>
    '''

@app.route('/convert', methods=['POST'])
def convert_file():
    if 'file' not in request.files:
        return jsonify({'success': False, 'error': '파일이 선택되지 않았습니다.'}), 400
    
    file = request.files['file']
    if file.filename == '':
        return jsonify({'success': False, 'error': '파일이 선택되지 않았습니다.'}), 400
    
    # 파일명 안전하게 처리
    original_filename = file.filename
    safe_name = safe_filename(original_filename)
    
    # 확장자 확인
    if not safe_name.lower().endswith('.pdf'):
        return jsonify({'success': False, 'error': 'PDF 파일만 업로드 가능합니다.'}), 400
    
    try:
        # 임시 파일로 저장
        with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as temp_pdf:
            file.save(temp_pdf.name)
            temp_pdf_path = temp_pdf.name
        
        # 출력 파일명 생성
        base_name = safe_name.rsplit('.', 1)[0]
        output_filename = f"{base_name}_텍스트변환.docx"
        
        # outputs 폴더 확인/생성
        output_dir = "outputs"
        if not os.path.exists(output_dir):
            os.makedirs(output_dir)
        
        output_path = os.path.join(output_dir, output_filename)
        
        # PDF → DOCX 변환 (텍스트만)
        if pdf_to_docx_text_only(temp_pdf_path, output_path):
            return jsonify({'success': True, 'filename': output_filename})
        else:
            return jsonify({'success': False, 'error': '변환 중 오류가 발생했습니다.'}), 500
            
    except Exception as e:
        print(f"변환 오류: {e}")
        return jsonify({'success': False, 'error': str(e)}), 500
    finally:
        # 임시 파일 삭제
        if 'temp_pdf_path' in locals() and os.path.exists(temp_pdf_path):
            try:
                os.unlink(temp_pdf_path)
            except:
                pass

@app.route('/download/<filename>')
def download_file(filename):
    try:
        output_path = os.path.join("outputs", filename)
        if os.path.exists(output_path):
            return send_file(output_path, as_attachment=True)
        else:
            return "파일을 찾을 수 없습니다.", 404
    except Exception as e:
        return f"다운로드 오류: {e}", 500

if __name__ == '__main__':
    print("🚀 PDF 텍스트 변환 서버 시작...")
    print("📄 원본 이미지 없이 편집 가능한 텍스트만 추출합니다!")
    app.run(debug=True, host='127.0.0.1', port=5000)