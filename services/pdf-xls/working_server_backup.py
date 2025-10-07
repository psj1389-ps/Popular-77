from flask import Flask, request, render_template, send_file, jsonify
import os
import tempfile
import subprocess
import platform
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.shared import Mm
from docx.oxml.shared import OxmlElement, qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024  # 100MB

# 폴더 생성
UPLOAD_FOLDER = 'uploads'
OUTPUT_FOLDER = 'outputs'
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

def allowed_file(filename, content_type=None):
    """파일 형식 확인 (확장자 또는 MIME 타입 기반)"""
    # 확장자로 확인
    if '.' in filename:
        ext = filename.rsplit('.', 1)[1].lower()
        return ext in {'pdf', 'docx'}
    
    # MIME 타입으로 확인 (확장자가 없는 경우)
    if content_type:
        return ('pdf' in content_type or 
                'document' in content_type or 
                'word' in content_type)
    
    return False

def setup_korean_fonts():
    """한글 폰트 설정"""
    try:
        # Windows 시스템 폰트 경로들
        font_paths = [
            r"C:\Windows\Fonts\malgun.ttf",  # 맑은 고딕
            r"C:\Windows\Fonts\gulim.ttc",   # 굴림
            r"C:\Windows\Fonts\batang.ttc",  # 바탕
            r"C:\Windows\Fonts\dotum.ttc",   # 돋움
        ]
        
        for font_path in font_paths:
            if os.path.exists(font_path):
                try:
                    pdfmetrics.registerFont(TTFont('Korean', font_path))
                    print(f"한글 폰트 등록 성공: {font_path}")
                    return True
                except Exception as e:
                    print(f"폰트 등록 실패: {font_path} - {e}")
                    continue
        
        print("한글 폰트를 찾을 수 없습니다. 기본 폰트를 사용합니다.")
        return False
    except Exception as e:
        print(f"폰트 설정 오류: {e}")
        return False

def add_korean_text_to_docx(doc, text, font_name="맑은 고딕"):
    """DOCX에 한글 텍스트 추가"""
    try:
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(text)
        
        # 한글 폰트 설정 시도
        try:
            run.font.name = font_name
            run.font.size = Pt(11)
        except:
            # 폰트 설정 실패 시 기본 설정 사용
            run.font.size = Pt(11)
        
        return True
    except Exception as e:
        print(f"한글 텍스트 추가 오류: {e}")
        return False

def detect_image_orientation(image):
    """이미지 방향 감지 (가로/세로) - A4 표준 비율 기준"""
    try:
        width, height = image.size
        aspect_ratio = width / height
        
        # A4 표준 비율 (297/210 ≈ 1.414)
        a4_landscape_ratio = 297 / 210  # 가로형 A4 비율
        a4_portrait_ratio = 210 / 297   # 세로형 A4 비율
        
        print(f"  - 이미지 크기: {width} x {height} 픽셀")
        print(f"  - 이미지 비율: {aspect_ratio:.3f}")
        print(f"  - A4 가로형 비율: {a4_landscape_ratio:.3f}")
        print(f"  - A4 세로형 비율: {a4_portrait_ratio:.3f}")
        
        # A4 표준 비율과 비교하여 방향 판단
        if aspect_ratio >= 1.3:  # A4 가로형에 가까운 비율
            print(f"  - ✅ 가로형 감지됨 (비율: {aspect_ratio:.3f} >= 1.3)")
            return "landscape"
        elif aspect_ratio <= 0.8:  # A4 세로형에 가까운 비율
            print(f"  - ✅ 세로형 감지됨 (비율: {aspect_ratio:.3f} <= 0.8)")
            return "portrait"
        else:  # 중간 비율인 경우 - A4 표준과 더 가까운 쪽으로 판단
            landscape_diff = abs(aspect_ratio - a4_landscape_ratio)
            portrait_diff = abs(aspect_ratio - a4_portrait_ratio)
            
            if landscape_diff < portrait_diff:
                print(f"  - ✅ A4 가로형에 더 가까움, 가로형으로 처리 (비율: {aspect_ratio:.3f})")
                return "landscape"
            else:
                print(f"  - ✅ A4 세로형에 더 가까움, 세로형으로 처리 (비율: {aspect_ratio:.3f})")
                return "portrait"
    except Exception as e:
        print(f"이미지 방향 감지 오류: {e}")
        return "landscape"  # 기본값을 가로로 변경

def get_image_size_for_docx(image, orientation):
    """화면 비율에 맞게 A4 표준 크기로 조정 (가로형: 297X210, 세로형: 210X297)"""
    try:
        width, height = image.size
        print(f"  - 원본 이미지 크기: {width} x {height} 픽셀")
        
        # 정확한 비율 계산
        aspect_ratio = width / height
        print(f"  - 이미지 비율 (가로/세로): {aspect_ratio:.3f}")
        
        if orientation == "landscape":
            # 가로 문서: 297mm X 210mm (A4 가로 모드)
            # 297mm를 인치로 변환 (25.4mm = 1인치)
            target_width = 297 / 25.4  # 약 11.69인치
            target_height = 210 / 25.4  # 약 8.27인치
            
            print(f"  - 🖼️ 가로형 A4 크기: {target_width:.2f} x {target_height:.2f} 인치 (297mm x 210mm)")
            print(f"  - 가로형 확인: 가로({target_width:.2f}) > 세로({target_height:.2f})")
            return Inches(target_width), Inches(target_height)
            
        else:
            # 세로 문서: 210mm X 297mm (A4 세로 모드)
            # 210mm를 인치로 변환
            target_width = 210 / 25.4  # 약 8.27인치
            target_height = 297 / 25.4  # 약 11.69인치
            
            print(f"  - 📄 세로형 A4 크기: {target_width:.2f} x {target_height:.2f} 인치 (210mm x 297mm)")
            print(f"  - 세로형 확인: 세로({target_height:.2f}) > 가로({target_width:.2f})")
            return Inches(target_width), Inches(target_height)
            
    except Exception as e:
        print(f"이미지 크기 계산 오류: {e}")
        # 기본값: A4 세로 크기
        return Inches(210 / 25.4), Inches(297 / 25.4)

def _set_section_orientation(section, orientation: str):
    """섹션 용지 방향과 크기를 A4에 맞춰 설정한다."""
    if orientation == "landscape":
        section.orientation = WD_ORIENT.LANDSCAPE
        section.page_width = Mm(297)
        section.page_height = Mm(210)
    else:
        section.orientation = WD_ORIENT.PORTRAIT
        section.page_width = Mm(210)
        section.page_height = Mm(297)
    margin = Mm(15)
    section.top_margin = margin
    section.bottom_margin = margin
    section.left_margin = margin
    section.right_margin = margin


def _fit_dimensions_within(max_width_inch: float, max_height_inch: float, img_width_inch: float, img_height_inch: float):
    """주어진 영역 안에 비율을 유지하며 이미지를 맞춘다."""
    scale = min(max_width_inch / img_width_inch, max_height_inch / img_height_inch)
    return Inches(img_width_inch * scale), Inches(img_height_inch * scale)


def pdf_to_docx(pdf_path, output_path):
    """PDF를 DOCX로 변환 (한글 폰트 지원, 페이지별 방향/크기 적용)"""
    try:
        print(f"PDF → DOCX 변환 시작: {pdf_path}")
        
        # PDF를 이미지로 변환
        images = convert_from_path(pdf_path, dpi=200)
        
        # 새 Word 문서 생성
        doc = Document()
        
        # 변환 정보 없이 바로 원본 내용만 표시
        
        for i, image in enumerate(images):
            print(f"페이지 {i+1}/{len(images)} 처리 중...")
            
            # 이미지 방향 감지
            orientation = detect_image_orientation(image)
            print(f"  - 이미지 방향: {orientation}")
            
            # 섹션 방향/용지 크기 설정
            if i == 0:
                section = doc.sections[0]
            else:
                section = doc.add_section(WD_SECTION.NEW_PAGE)
            _set_section_orientation(section, orientation)

            # 본문 최대 영역 계산 (여백 제외)
            max_w_inch = float(section.page_width.inches - (section.left_margin.inches + section.right_margin.inches))
            max_h_inch = float(section.page_height.inches - (section.top_margin.inches + section.bottom_margin.inches))

            # 원본 이미지 크기(inch) 추정 - 200DPI 기준
            dpi = 200
            img_w_inch = image.size[0] / dpi
            img_h_inch = image.size[1] / dpi

            fitted_w, fitted_h = _fit_dimensions_within(max_w_inch, max_h_inch, img_w_inch, img_h_inch)

            # 원본 이미지 추가
            print("  - 이미지 추가 중...")
            temp_file_path = None
            try:
            with tempfile.NamedTemporaryFile(suffix='.jpg', delete=False) as temp_file:
                    temp_file_path = temp_file.name
                    image.save(temp_file_path, 'JPEG', quality=90, optimize=True)

                doc.add_picture(temp_file_path, width=fitted_w, height=fitted_h)
                
            finally:
                # 임시 파일 삭제 (안전하게)
                if temp_file_path and os.path.exists(temp_file_path):
                    try:
                        os.unlink(temp_file_path)
                    except (OSError, PermissionError) as e:
                        print(f"임시 파일 삭제 실패 (무시됨): {e}")
                        # 임시 파일 삭제 실패는 무시하고 계속 진행
            
            # 섹션으로 페이지가 구분되므로 page_break는 불필요
        
        # DOCX 파일 저장
        doc.save(output_path)
        print(f"✅ PDF → DOCX 변환 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ PDF → DOCX 변환 오류: {e}")
        import traceback
        traceback.print_exc()
        return False

def docx_to_pdf(docx_path, output_path):
    """DOCX를 PDF로 변환 (한글 폰트 지원)"""
    try:
        print(f"DOCX → PDF 변환 시작: {docx_path}")
        
        # 한글 폰트 설정
        font_setup = setup_korean_fonts()
        
        # DOCX 문서 읽기
        doc = Document(docx_path)
        
        # PDF 생성
        c = canvas.Canvas(output_path, pagesize=A4)
        width, height = A4
        y_position = height - 50
        
        # 한글 폰트 사용
        if font_setup:
            c.setFont("Korean", 12)
        else:
            c.setFont("Helvetica", 12)
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                # 텍스트 처리 (한글 지원)
                text = paragraph.text.strip()
                
                # 긴 텍스트를 여러 줄로 나누기
                max_chars_per_line = 80
                lines = []
                while len(text) > max_chars_per_line:
                    lines.append(text[:max_chars_per_line])
                    text = text[max_chars_per_line:]
                if text:
                    lines.append(text)
                
                # 각 줄을 PDF에 추가
                for line in lines:
                if y_position < 50:
                    c.showPage()
                    y_position = height - 50
                        if font_setup:
                            c.setFont("Korean", 12)
                        else:
                            c.setFont("Helvetica", 12)
                    
                    c.drawString(50, y_position, line)
                    y_position -= 20
        
        c.save()
        print(f"✅ DOCX → PDF 변환 완료: {output_path}")
        return True
        
    except Exception as e:
        print(f"❌ DOCX → PDF 변환 오류: {e}")
        import traceback
        traceback.print_exc()
        return False

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/convert', methods=['POST'])
def convert_file():
    try:
        print("=== 변환 요청 시작 ===")
        
        # 1. 파일 확인
        if 'file' not in request.files:
            print("오류: 파일이 없음")
            return jsonify({'success': False, 'error': '파일이 선택되지 않았습니다.'}), 400
        
        file = request.files['file']
        if file.filename == '':
            print("오류: 파일명이 없음")
            return jsonify({'success': False, 'error': '파일명이 비어있습니다.'}), 400
        
        # 2. 파일 형식 확인
        if not allowed_file(file.filename, file.content_type):
            print(f"오류: 지원하지 않는 파일 형식 - {file.filename} (MIME: {file.content_type})")
            return jsonify({'success': False, 'error': 'PDF 또는 DOCX 파일만 업로드 가능합니다.'}), 400
        
        # 3. 파일명 처리 및 저장
        original_filename = file.filename
        print(f"원본 파일명: {original_filename}")
        
        # 파일명 정리 및 처리
        cleaned_filename = original_filename.strip()
        if not cleaned_filename:
            cleaned_filename = "uploaded_file"
        
        # 파일명과 확장자를 분리하여 안전하게 처리
        if '.' in cleaned_filename:
            # 확장자가 있는 경우
            name_part, ext_part = cleaned_filename.rsplit('.', 1)
            safe_name = secure_filename(name_part) or "file"
            safe_ext = ext_part.lower().strip()
            
            # 확장자가 비어있으면 기본값 설정
            if not safe_ext:
                safe_ext = "pdf"
            
            filename = f"{safe_name}.{safe_ext}"
            print(f"확장자 분리 처리: {cleaned_filename} → {filename}")
        else:
            # 확장자가 없는 경우 MIME 타입으로 추정
            content_type = file.content_type
            print(f"파일 MIME 타입: {content_type}")
            
            safe_name = secure_filename(cleaned_filename) or "file"
            
            if 'pdf' in content_type:
                filename = f"{safe_name}.pdf"
                print(f"PDF 파일로 추정하여 .pdf 확장자 추가")
            elif 'document' in content_type or 'word' in content_type:
                filename = f"{safe_name}.docx"
                print(f"DOCX 파일로 추정하여 .docx 확장자 추가")
            else:
                # MIME 타입도 없는 경우 기본값으로 처리
                print(f"경고: 파일에 확장자와 MIME 타입이 없음 - {cleaned_filename}")
                filename = f"{safe_name}.pdf"
                print(f"기본값으로 .pdf 확장자 추가")
        
        # 최종 파일명 검증 및 보정
        if not filename or filename == '.' or filename == '..' or '.' not in filename:
            filename = "uploaded_file.pdf"
            print(f"안전하지 않은 파일명으로 인해 기본 파일명 사용: {filename}")
        
        # 확장자 최종 검증
        if not filename.endswith(('.pdf', '.docx')):
            if filename.endswith('.pdf') or 'pdf' in file.content_type:
                filename = filename.rsplit('.', 1)[0] + '.pdf'
            else:
                filename = filename.rsplit('.', 1)[0] + '.docx'
            print(f"확장자 보정: {filename}")
        
        file_path = os.path.join(UPLOAD_FOLDER, filename)
        file.save(file_path)
        print(f"파일 저장 완료: {file_path}")
        
        # 4. 파일 확장자 확인 (간단하고 안전한 처리)
        print(f"최종 파일명: {filename}")
        
        # 확장자 추출 (이미 위에서 안전하게 처리되었으므로 간단하게)
        if '.' not in filename:
            print(f"오류: 최종 파일명에 확장자가 없음 - {filename}")
            return jsonify({'success': False, 'error': '파일 확장자를 확인할 수 없습니다.'}), 400
        
        try:
            file_ext = filename.split('.')[-1].lower().strip()
            if not file_ext:
                print(f"오류: 확장자가 비어있음 - {filename}")
                return jsonify({'success': False, 'error': '파일 확장자를 확인할 수 없습니다.'}), 400
                
        print(f"파일 확장자: {file_ext}")
            
        except Exception as e:
            print(f"오류: 파일 확장자 추출 중 예외 발생 - {e}")
            return jsonify({'success': False, 'error': '파일 확장자를 확인할 수 없습니다.'}), 400
        
        # 5. 변환 처리
        try:
            # 출력 파일명 생성 (안전하게)
            base_name = filename.rsplit('.', 1)[0].strip()
            if not base_name:
                base_name = "converted_file"  # 기본 파일명
            
        if file_ext == 'pdf':
            # PDF → DOCX
                output_filename = base_name + '.docx'
            output_path = os.path.join(OUTPUT_FOLDER, output_filename)
            print(f"PDF → DOCX 변환: {file_path} → {output_path}")
            success = pdf_to_docx(file_path, output_path)
            
        elif file_ext == 'docx':
            # DOCX → PDF
                output_filename = base_name + '.pdf'
            output_path = os.path.join(OUTPUT_FOLDER, output_filename)
            print(f"DOCX → PDF 변환: {file_path} → {output_path}")
            success = docx_to_pdf(file_path, output_path)
            else:
                print(f"오류: 지원하지 않는 파일 형식 - {file_ext}")
                return jsonify({'success': False, 'error': f'지원하지 않는 파일 형식입니다: {file_ext}'}), 400
                
        except Exception as e:
            print(f"오류: 변환 처리 중 예외 발생 - {e}")
            return jsonify({'success': False, 'error': '파일 변환 처리 중 오류가 발생했습니다.'}), 500
        
        # 6. 임시 파일 삭제
        try:
            os.remove(file_path)
            print("임시 파일 삭제 완료")
        except:
            pass
        
        # 7. 결과 처리
        if success and os.path.exists(output_path):
            print(f"✅ 변환 성공! 다운로드: {output_filename}")
            return send_file(output_path, as_attachment=True, download_name=output_filename)
        else:
            print("❌ 변환 실패")
            return jsonify({'success': False, 'error': '파일 변환에 실패했습니다.'}), 500
            
    except Exception as e:
        print(f"❌ 서버 오류: {e}")
        import traceback
        traceback.print_exc()
        return jsonify({'success': False, 'error': f'서버 오류: {str(e)}'}), 500

@app.errorhandler(413)
def too_large(e):
    return jsonify({'success': False, 'error': '파일 크기가 100MB를 초과합니다.'}), 413

if __name__ == '__main__':
    print("🚀 PDF ↔ DOCX 변환기 시작")
    print("📍 서버 주소: http://127.0.0.1:5000")
    print("📍 네트워크: http://0.0.0.0:5000")
    
    # 한글 폰트 설정
    print("🔤 한글 폰트 설정 중...")
    font_setup = setup_korean_fonts()
    if font_setup:
        print("✅ 한글 폰트 설정 완료")
    else:
        print("⚠️ 한글 폰트 설정 실패 - 기본 폰트 사용")
    
    app.run(debug=True, host='0.0.0.0', port=5000)