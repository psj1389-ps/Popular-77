from flask import Flask, render_template, request, jsonify, send_file
import os
import time
from werkzeug.utils import secure_filename
from pdf2image import convert_from_path
from docx import Document
from docx.shared import Inches
from docx.enum.section import WD_ORIENT
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import A4, landscape, portrait
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.lib.utils import ImageReader
import urllib.request
import PyPDF2
import unicodedata
import sys
from PIL import Image as PILImage
import io

# OCR 기능 확인 및 설정
try:
    import pytesseract
    import cv2
    import numpy as np
    OCR_AVAILABLE = True
    print("✅ OCR 모듈 로드 성공")
except ImportError:
    OCR_AVAILABLE = False
    print("⚠️ OCR 모듈 없음")

app = Flask(__name__)
app.config['MAX_CONTENT_LENGTH'] = 100 * 1024 * 1024

# 폴더 생성
os.makedirs('uploads', exist_ok=True)
os.makedirs('outputs', exist_ok=True)
os.makedirs('fonts', exist_ok=True)

# 한글 폰트 설정
KOREAN_FONT = 'Helvetica'
KOREAN_FONT_AVAILABLE = False
AVAILABLE_FONTS = {}  # 추가된 변수 정의

def setup_korean_font_advanced():
    """고급 한글 폰트 설정"""
    global KOREAN_FONT, KOREAN_FONT_AVAILABLE, AVAILABLE_FONTS
    
    # 1. 나눔고딕 TTF 시도
    try:
        font_path = os.path.join('fonts', 'NanumGothic.ttf')
        
        if not os.path.exists(font_path):
            print("📥 나눔고딕 폰트 다운로드 중...")
            font_url = "https://github.com/naver/nanumfont/raw/master/TTF/NanumGothic.ttf"
            
            req = urllib.request.Request(font_url, headers={
                'User-Agent': 'Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36'
            })
            
            with urllib.request.urlopen(req, timeout=30) as response:
                with open(font_path, 'wb') as f:
                    f.write(response.read())
        
        if os.path.exists(font_path):
            pdfmetrics.registerFont(TTFont('NanumGothic', font_path))
            KOREAN_FONT = 'NanumGothic'
            KOREAN_FONT_AVAILABLE = True
            AVAILABLE_FONTS['NanumGothic'] = {
                'path': font_path,
                'display_name': '나눔고딕'
            }
            print("✅ 나눔고딕 TTF 폰트 등록 완료")
            return True
            
    except Exception as e:
        print(f"나눔고딕 TTF 등록 실패: {e}")
    
    # 2. 시스템 한글 폰트 시도
    system_fonts = [
        (r'C:\Windows\Fonts\malgun.ttf', 'Malgun', '맑은 고딕'),
        (r'C:\Windows\Fonts\gulim.ttc', 'Gulim', '굴림'),
        (r'C:\Windows\Fonts\batang.ttc', 'Batang', '바탕'),
    ]
    
    for font_path, font_name, display_name in system_fonts:
        try:
            if os.path.exists(font_path):
                pdfmetrics.registerFont(TTFont(font_name, font_path))
                KOREAN_FONT = font_name
                KOREAN_FONT_AVAILABLE = True
                AVAILABLE_FONTS[font_name] = {
                    'path': font_path,
                    'display_name': display_name
                }
                print(f"✅ 시스템 한글 폰트 등록: {display_name}")
                return True
        except Exception as e:
            continue
    
    print("⚠️ 한글 폰트 등록 실패, 기본 폰트 사용")
    return False

# 앱 시작 시 한글 폰트 설정
# 기존 setup_korean_font_advanced() 함수를 setup_korean_font_simple()로 교체
# 온라인 폰트 다운로드 코드 제거
# 시스템 폰트만 사용
setup_korean_font_advanced()

def safe_korean_text(text):
    """한글 텍스트 안전 처리"""
    if not text:
        return ""
    
    try:
        normalized = unicodedata.normalize('NFC', str(text))
        cleaned = ''.join(char for char in normalized 
                         if unicodedata.category(char) not in ['Cc', 'Cf'])
        return cleaned if cleaned.strip() else ""
    except Exception as e:
        return str(text) if text else ""

def draw_korean_text(canvas_obj, x, y, text, font_size=11):
    """한글 텍스트 그리기"""
    if not text or not text.strip():
        return
    
    try:
        safe_text = safe_korean_text(text)
        if not safe_text:
            return
        
        if KOREAN_FONT_AVAILABLE:
            canvas_obj.setFont(KOREAN_FONT, font_size)
            canvas_obj.drawString(x, y, safe_text)
        else:
            canvas_obj.setFont('Helvetica', font_size)
            # 한글이 있으면 대체 문자 사용
            has_korean = any('\uac00' <= char <= '\ud7af' for char in safe_text)
            if has_korean:
                display_text = ''.join('한' if '\uac00' <= char <= '\ud7af' else char for char in safe_text)
            else:
                display_text = safe_text
            canvas_obj.drawString(x, y, display_text)
            
    except Exception as e:
        try:
            canvas_obj.setFont('Helvetica', 8)
            canvas_obj.drawString(x, y, "[Error]")
        except:
            pass

def extract_images_from_docx(docx_path, temp_files):
    """DOCX에서 모든 이미지 추출 (강화된 버전)"""
    images = []
    
    try:
        print("🖼️ DOCX에서 이미지 추출 시작...")
        
        # DOCX 파일을 ZIP으로 열어서 이미지 직접 추출
        with zipfile.ZipFile(docx_path, 'r') as docx_zip:
            # media 폴더의 모든 이미지 파일 찾기
            media_files = [f for f in docx_zip.namelist() if f.startswith('word/media/')]
            
            for i, media_file in enumerate(media_files):
                try:
                    # 이미지 파일 확장자 확인
                    if any(media_file.lower().endswith(ext) for ext in ['.jpg', '.jpeg', '.png', '.gif', '.bmp']):
                        # 이미지 데이터 추출
                        image_data = docx_zip.read(media_file)
                        
                        # PIL로 이미지 정보 확인
                        pil_image = PILImage.open(io.BytesIO(image_data))
                        width, height = pil_image.size
                        
                        # 임시 파일로 저장
                        timestamp = str(int(time.time() * 1000))
                        temp_img_path = os.path.join('uploads', f'extracted_img_{timestamp}_{i}.jpg')
                        
                        # JPEG로 변환하여 저장
                        if pil_image.mode in ('RGBA', 'LA', 'P'):
                            # 투명도가 있는 이미지는 흰 배경으로 변환
                            background = PILImage.new('RGB', pil_image.size, (255, 255, 255))
                            if pil_image.mode == 'P':
                                pil_image = pil_image.convert('RGBA')
                            background.paste(pil_image, mask=pil_image.split()[-1] if pil_image.mode == 'RGBA' else None)
                            pil_image = background
                        
                        pil_image.save(temp_img_path, 'JPEG', quality=90)
                        temp_files.append(temp_img_path)
                        
                        images.append({
                            'path': temp_img_path,
                            'width': width,
                            'height': height,
                            'original_name': media_file
                        })
                        
                        print(f"✅ 이미지 추출: {media_file} ({width}x{height})")
                        
                except Exception as e:
                    print(f"이미지 {media_file} 추출 오류: {e}")
                    continue
        
        print(f"✅ 총 {len(images)}개 이미지 추출 완료")
        return images
        
    except Exception as e:
        print(f"❌ 이미지 추출 실패: {e}")
        return []

import zipfile

def extract_docx_with_complete_formatting(docx_path, temp_files):
    """DOCX에서 완전한 서식 정보와 함께 내용 추출"""
    try:
        doc = Document(docx_path)
        all_content = []
        
        print("📝 완전한 서식 정보와 함께 내용 추출 시작...")
        
        # 1. 이미지 먼저 추출
        extracted_images = extract_images_from_docx(docx_path, temp_files)
        image_index = 0
        
        # 2. 문단별 서식 정보 추출 (이미지 포함)
        for i, paragraph in enumerate(doc.paragraphs):
            try:
                # 문단에 이미지가 있는지 확인
                has_image = False
                for run in paragraph.runs:
                    if hasattr(run, '_element'):
                        # drawing 요소 확인 (이미지)
                        drawings = run._element.xpath('.//a:blip')
                        if drawings and image_index < len(extracted_images):
                            # 이미지 추가
                            img_info = extracted_images[image_index]
                            all_content.append({
                                'type': 'image',
                                'path': img_info['path'],
                                'width': img_info['width'],
                                'height': img_info['height'],
                                'ocr_text': img_info.get('ocr_text', ''),
                                'index': image_index
                            })
                            print(f"📷 이미지 {image_index + 1} 위치 확인: {img_info['original_name']}")
                            image_index += 1
                            has_image = True
                
                # 250-290번째 줄 영역 수정
                # 텍스트 처리 (향상된 서식 추출)
                if paragraph.text and paragraph.text.strip():
                    text = safe_korean_text(paragraph.text.strip())
                    
                    # 향상된 서식 정보 추출
                    formatting = extract_enhanced_formatting(paragraph)
                    
                    if text:  # 텍스트가 있을 때만 추가
                        all_content.append({
                            'type': 'paragraph',
                            'content': text,
                            'formatting': formatting,
                            'index': i
                        })
                        
                        print(f"문단 {i+1}: {text[:20]}... (크기: {formatting['font_size']}, 굵게: {formatting['is_bold']}, 색상: {formatting['color']})")
                    
            except Exception as e:
                print(f"문단 {i} 처리 오류: {e}")
                continue
        
        # 3. 남은 이미지들 추가 (문단에 포함되지 않은 이미지)
        while image_index < len(extracted_images):
            img_info = extracted_images[image_index]
            all_content.append({
                'type': 'image',
                'path': img_info['path'],
                'width': img_info['width'],
                'height': img_info['height'],
                'ocr_text': img_info.get('ocr_text', ''),
                'index': image_index
            })
            print(f"📷 추가 이미지 {image_index + 1}: {img_info['original_name']}")
            image_index += 1
        
        # 4. 표 추출 (서식 포함)
        for table_idx, table in enumerate(doc.tables):
            try:
                table_content = []
                for row_idx, row in enumerate(table.rows):
                    row_content = []
                    for cell_idx, cell in enumerate(row.cells):
                        try:
                            cell_text = safe_korean_text(cell.text.strip())
                            if cell_text:
                                row_content.append(cell_text)
                        except:
                            row_content.append("")
                    
                    if row_content and any(row_content):
                        table_content.append(row_content)
                
                if table_content:
                    all_content.append({
                        'type': 'table',
                        'content': table_content,
                        'index': table_idx
                    })
                    print(f"표 {table_idx+1}: {len(table_content)}행")
            except Exception as e:
                print(f"표 {table_idx} 처리 오류: {e}")
                continue
        
        print(f"✅ 총 {len(all_content)}개 요소 추출 (이미지 {len(extracted_images)}개 포함)")
        return all_content
        
    except Exception as e:
        print(f"❌ DOCX 완전 서식 추출 실패: {e}")
        return []

def detect_pdf_orientation(pdf_path):
    """PDF 문서의 방향 감지"""
    try:
        with open(pdf_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            
            if len(pdf_reader.pages) > 0:
                first_page = pdf_reader.pages[0]
                
                if hasattr(first_page, 'mediabox'):
                    mediabox = first_page.mediabox
                    width = float(mediabox.width)
                    height = float(mediabox.height)
                    
                    rotation = 0
                    if hasattr(first_page, 'rotation'):
                        rotation = first_page.rotation or 0
                    
                    if rotation in [90, 270]:
                        width, height = height, width
                    
                    if width > height:
                        return 'landscape', width, height
                    else:
                        return 'portrait', width, height
        
        return 'portrait', 595, 842
        
    except Exception as e:
        print(f"⚠️ PDF 방향 감지 실패: {e}")
        return 'portrait', 595, 842

def detect_docx_orientation(docx_path):
    """DOCX 문서의 방향 감지"""
    try:
        doc = Document(docx_path)
        
        if doc.sections and len(doc.sections) > 0:
            section = doc.sections[0]
            if hasattr(section, 'page_width') and hasattr(section, 'page_height'):
                width = section.page_width.inches
                height = section.page_height.inches
                
                if width > height:
                    return 'landscape'
                else:
                    return 'portrait'
        
        return 'portrait'
        
    except Exception as e:
        print(f"⚠️ DOCX 방향 감지 실패: {e}")
        return 'portrait'

def set_docx_orientation(doc, orientation):
    """DOCX 문서의 방향 설정"""
    try:
        for section in doc.sections:
            if orientation == 'landscape':
                section.orientation = WD_ORIENT.LANDSCAPE
                section.page_width = Inches(11.69)
                section.page_height = Inches(8.27)
            else:
                section.orientation = WD_ORIENT.PORTRAIT
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
        return True
    except Exception as e:
        print(f"⚠️ DOCX 방향 설정 실패: {e}")
        return False

def safe_file_check(filename):
    """파일 확장자 안전 확인"""
    try:
        if not filename or '.' not in filename:
            return False, 'unknown'
        
        extension = filename.lower().split('.')[-1]
        
        if extension in ['pdf', 'docx']:
            return True, extension
        else:
            return False, extension
            
    except Exception as e:
        print(f"파일 확인 오류: {e}")
        return False, 'unknown'

def clean_temp_files(file_list):
    """임시 파일 안전 삭제"""
    for file_path in file_list:
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except:
            pass

@app.route('/')
def index():
    return render_template('index.html')

@app.route('/fonts')
def list_fonts():
    """사용 가능한 폰트 목록 API"""
    font_list = []
    for font_name, font_info in AVAILABLE_FONTS.items():
        font_list.append({
            'name': font_name,
            'display_name': font_info['display_name'],
            'current': font_name == KOREAN_FONT
        })
    
    return jsonify({
        'fonts': font_list,
        'current_font': KOREAN_FONT,
        'total_fonts': len(AVAILABLE_FONTS)
    })

@app.route('/convert', methods=['POST'])
def convert_file():
    temp_files = []
    
    try:
        print("=== PDF ↔ DOCX 변환 시작 ===")
        print(f"🔤 사용 가능한 폰트: {len(AVAILABLE_FONTS)}개")
        # OCR 관련 출력 제거
        
        # 1. 파일 확인
        if 'file' not in request.files:
            return jsonify({'success': False, 'error': '파일이 선택되지 않았습니다.'}), 400
        
        file = request.files['file']
        if not file.filename:
            return jsonify({'success': False, 'error': '파일명이 없습니다.'}), 400
        
        print(f"업로드된 파일: {file.filename}")
        
        # 2. 파일 형식 확인
        is_valid, extension = safe_file_check(file.filename)
        if not is_valid:
            return jsonify({
                'success': False, 
                'error': f'지원하지 않는 파일 형식입니다. PDF 또는 DOCX 파일만 업로드 가능합니다. (현재: {extension})'
            }), 400
        
        # 3. 파일 저장
        timestamp = str(int(time.time() * 1000))
        safe_filename = secure_filename(file.filename)
        name_without_ext = safe_filename.rsplit('.', 1)[0] if '.' in safe_filename else safe_filename
        input_path = os.path.join('uploads', f"{name_without_ext}_{timestamp}.{extension}")
        
        file.save(input_path)
        temp_files.append(input_path)
        print(f"✅ 파일 저장: {input_path}")
        
        # 4. 변환 처리
        if extension == 'pdf':
            # PDF → DOCX
            output_path = os.path.join('outputs', f"{name_without_ext}_{timestamp}.docx")
            
            try:
                print("📄 PDF → DOCX 변환 시작")
                
                pdf_orientation, pdf_width, pdf_height = detect_pdf_orientation(input_path)
                images = convert_from_path(input_path, dpi=150)
                
                doc = Document()
                set_docx_orientation(doc, pdf_orientation)
                
                success_count = 0
                for i, img in enumerate(images):
                    try:
                        img_path = os.path.join('uploads', f'page_{timestamp}_{i}.jpg')
                        temp_files.append(img_path)
                        
                        img.save(img_path, 'JPEG', quality=85)
                        
                        if pdf_orientation == 'landscape':
                            doc.add_picture(img_path, width=Inches(9))
                        else:
                            doc.add_picture(img_path, width=Inches(6))
                        
                        if i < len(images) - 1:
                            doc.add_page_break()
                        
                        success_count += 1
                        
                    except Exception as e:
                        print(f"⚠️ 페이지 {i+1} 처리 오류: {e}")
                        continue
                
                if success_count == 0:
                    doc.add_paragraph("PDF 변환 완료")
                    doc.add_paragraph(f"원본 파일: {file.filename}")
                
                doc.save(output_path)
                print(f"✅ DOCX 저장 완료: {success_count}개 페이지")
                
            except Exception as e:
                print(f"❌ PDF 변환 오류: {e}")
                doc = Document()
                doc.add_paragraph("PDF 변환 중 오류 발생")
                doc.save(output_path)
        
        elif extension == 'docx':
            # DOCX → PDF (다중 폰트 지원)
            output_path = os.path.join('outputs', f"{name_without_ext}_{timestamp}.pdf")
            
            try:
                print("📄 DOCX → PDF 변환 시작 (다중 폰트 지원)")
                
                # 방향 감지
                docx_orientation = detect_docx_orientation(input_path)
                
                # 강화된 서식 정보와 함께 내용 추출 (이미지 포함)
                # 557번째 줄 확인
                # 변경 전: extract_docx_with_complete_formatting_enhanced
                # 변경 후: extract_docx_with_complete_formatting
                content_list = extract_docx_with_complete_formatting(input_path, temp_files)
                
                if not content_list:
                    return jsonify({'success': False, 'error': 'DOCX 파일에서 내용을 추출할 수 없습니다.'}), 400
                
                # PDF 페이지 크기 설정
                if docx_orientation == 'landscape':
                    page_size = landscape(A4)
                    base_font_size = 10
                    line_height_base = 16
                    max_chars_per_line = 70
                else:
                    page_size = portrait(A4)
                    base_font_size = 11
                    line_height_base = 18
                    max_chars_per_line = 50
                
                # PDF 생성
                c = canvas.Canvas(output_path, pagesize=page_size)
                width, height = page_size
                
                print(f"📄 PDF 생성: {width:.0f} x {height:.0f} ({docx_orientation})")
                
                margin_left = 50
                margin_right = width - 50
                margin_top = height - 50
                margin_bottom = 50
                
                y_pos = margin_top
                
                # 내용 처리 (다중 폰트 지원)
                processed_items = 0
                image_count = 0
                
                # 내용 처리 (576번째 줄 이후)
                for item in content_list:
                    try:
                        if item['type'] == 'paragraph':
                            text = item['content']
                            formatting = item.get('formatting', {})
                            
                            # 페이지 넘김 확인
                            if y_pos < margin_bottom + formatting.get('font_size', 11) * 2:
                                c.showPage()
                                y_pos = margin_top
                            
                            # 향상된 텍스트 그리기
                            draw_enhanced_text(c, margin_left, y_pos, text, formatting)
                            y_pos -= formatting.get('font_size', 11) + 5
                            
                        elif item['type'] == 'image':
                            # 페이지 넘김 확인
                            estimated_height = min(item.get('height', 200), 400)
                            if y_pos < margin_bottom + estimated_height:
                                c.showPage()
                                y_pos = margin_top
                            
                            # 향상된 이미지 그리기
                            image_height = draw_enhanced_image(
                                c, margin_left, y_pos, 
                                item['path'], 
                                item.get('width', 400), 
                                item.get('height', 300)
                            )
                            y_pos -= image_height
                            image_count += 1
                            
                        processed_items += 1
                        
                    except Exception as e:
                        print(f"항목 처리 오류: {e}")
                        continue
                
                # PDF 저장
                c.save()
                clean_temp_files(temp_files)
                
                print(f"✅ PDF 생성 완료: {processed_items}개 항목, {image_count}개 이미지")
                
                return send_file(output_path, as_attachment=True, 
                               download_name=f"{name_without_ext}.pdf")
                               
            except Exception as e:
                clean_temp_files(temp_files)
                return jsonify({'error': f'DOCX → PDF 변환 오류: {str(e)}'}), 500
        
        else:
            return jsonify({'error': '지원하지 않는 파일 형식입니다.'}), 400
            
    except Exception as e:
        clean_temp_files(temp_files)
        return jsonify({'error': f'파일 처리 오류: {str(e)}'}), 500
    
    finally:
        clean_temp_files(temp_files)
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left, y_pos, 
                            item['path'], 
                            item.get('width', 400), 
                            item.get('height', 300)
                        )
                        y_pos -= image_height
                        
                        # 향상된 이미지 그리기
                        image_height = draw_enhanced_image(
                            c, margin_left,